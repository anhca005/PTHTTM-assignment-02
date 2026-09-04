# -*- coding: utf-8 -*-
"""
Builds BaoCao_Assignment02.docx (Vietnamese) covering all three applications.
Run:  python build_report.py
Re-run any time after screenshots are added to report/images/<app>/<file>.png
(see SHOT_LIST.md) — placeholders are automatically replaced by the real
images once they exist at the expected path.
"""
import json
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = os.path.dirname(os.path.abspath(__file__))
A02 = os.path.dirname(ROOT)
IMG_DIR = os.path.join(ROOT, "images")
PLACEHOLDER_DIR = os.path.join(IMG_DIR, "_placeholders")
os.makedirs(PLACEHOLDER_DIR, exist_ok=True)

CHART = {
    "diabetes": os.path.join(A02, "diabetes", "charts"),
    "house_price": os.path.join(A02, "house_price", "charts"),
    "ecommerce": os.path.join(A02, "ecommerce", "charts"),
}


def load_json(app, name="meta.json"):
    with open(os.path.join(A02, app, "model", name), encoding="utf-8") as f:
        return json.load(f)


META = {
    "diabetes": load_json("diabetes"),
    "house_price": load_json("house_price"),
    "ecommerce": load_json("ecommerce"),
}

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Times New Roman")

for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)

FIG_COUNTER = {"n": 0}
TABLE_COUNTER = {"n": 0}


BLACK = RGBColor(0, 0, 0)


def add_heading(text, level=1):
    display_text = text.upper() if level == 1 else text
    h = doc.add_heading(display_text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = BLACK
        if level == 1:
            run.bold = True
    return h


def add_para(text="", bold=False, italic=False, size=13, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r.font.color.rgb = BLACK
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(13)
        r.font.name = "Times New Roman"
        r.font.color.rgb = BLACK


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(14)


def make_placeholder(app, filename, desc, size=(1400, 500)):
    """Creates a light-gray placeholder PNG with a dashed border and a
    caption describing which real screenshot should replace it."""
    out_path = os.path.join(PLACEHOLDER_DIR, f"{app}__{filename}")
    img = Image.new("RGB", size, color=(245, 246, 250))
    draw = ImageDraw.Draw(img)
    dash = 14
    for x in range(0, size[0], dash * 2):
        draw.line([(x, 4), (x + dash, 4)], fill=(150, 160, 190), width=3)
        draw.line([(x, size[1] - 4), (x + dash, size[1] - 4)], fill=(150, 160, 190), width=3)
    for y in range(0, size[1], dash * 2):
        draw.line([(4, y), (4, y + dash)], fill=(150, 160, 190), width=3)
        draw.line([(size[0] - 4, y), (size[0] - 4, y + dash)], fill=(150, 160, 190), width=3)
    try:
        font = ImageFont.truetype("arial.ttf", 26)
        font_small = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        font = ImageFont.load_default()
        font_small = font
    text1 = "[ ẢNH CHỤP TỪ JUPYTER SẼ ĐƯỢC CHÈN TẠI ĐÂY ]"
    text2 = desc
    text3 = f"file: images/{app}/{filename}"
    for i, (t, fnt) in enumerate([(text1, font), (text2, font_small), (text3, font_small)]):
        bbox = draw.textbbox((0, 0), t, font=fnt)
        w = bbox[2] - bbox[0]
        draw.text(((size[0] - w) / 2, size[1] / 2 - 50 + i * 36), t, fill=(80, 90, 130), font=fnt)
    img.save(out_path)
    return out_path


def add_figure(app, filename, desc, width_in=6.0):
    """Inserts the real screenshot if it exists at images/<app>/<filename>,
    otherwise inserts an auto-generated placeholder box."""
    real_path = os.path.join(IMG_DIR, app, filename)
    FIG_COUNTER["n"] += 1
    if os.path.exists(real_path):
        path = real_path
    else:
        path = make_placeholder(app, filename, desc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    add_caption(f"Hình {FIG_COUNTER['n']}. {desc}")


def add_table(headers, rows, widths=None):
    TABLE_COUNTER["n"] += 1
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        r.font.color.rgb = BLACK
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
            r.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_page_break():
    doc.add_page_break()


def fmt_pct(x):
    return f"{x*100:.2f}%"


def fmt_num(x, nd=4):
    return f"{x:.{nd}f}"


# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG")
r.bold = True
r.font.size = Pt(15)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("KHOA CÔNG NGHỆ THÔNG TIN 1")
r2.bold = True
r2.font.size = Pt(14)
for _ in range(4):
    doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("BÁO CÁO ASSIGNMENT 02")
r3.bold = True
r3.font.size = Pt(26)
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("MÔN HỌC: PHÁT TRIỂN CÁC HỆ THỐNG THÔNG MINH")
r4.bold = True
r4.font.size = Pt(15)
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("From Data Representation to a Deployable Intelligent System")
r5.italic = True
r5.font.size = Pt(13)
for _ in range(4):
    doc.add_paragraph()
info_lines = [
    "Giảng viên hướng dẫn : PGS.TS. Trần Đình Quế",
    "Sinh viên thực hiện    : ..................................",
    "Mã sinh viên            : ..................................",
    "Lớp/Nhóm học phần   : ..................................",
]
for line in info_lines:
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = pi.add_run(line)
    ri.bold = True
    ri.font.size = Pt(13)
for _ in range(6):
    doc.add_paragraph()
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("Hà Nội, 2026")
r6.italic = True
r6.font.size = Pt(13)
add_page_break()

# ---------------------------------------------------------------------------
# TABLE OF CONTENTS (Word field — press F9 / right-click "Update Field")
# ---------------------------------------------------------------------------
add_heading("Mục lục", level=1)
para = doc.add_paragraph()
run = para.add_run()
fld_begin = OxmlElement("w:fldChar")
fld_begin.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fld_sep = OxmlElement("w:fldChar")
fld_sep.set(qn("w:fldCharType"), "separate")
fld_text = OxmlElement("w:t")
fld_text.text = "Click chuột phải vào đây và chọn 'Update Field' để tạo Mục lục."
fld_end = OxmlElement("w:fldChar")
fld_end.set(qn("w:fldCharType"), "end")
run._r.append(fld_begin)
run._r.append(instr)
run._r.append(fld_sep)
run._r.append(fld_text)
run._r.append(fld_end)
add_page_break()

print("Header sections written. Continuing with topic-first sections 1-10...")

# Sections 1-10 follow the assignment's own required Report Structure
# (topic-first: one section per topic, spanning all three applications) —
# see sections.py. This is deliberately organized differently from a
# per-application chapter layout.
import sections  # noqa: E402

H = {
    "heading": add_heading, "para": add_para, "bullets": add_bullets,
    "table": add_table, "figure": add_figure,
    "page_break": add_page_break,
}

sections.section_1(doc, H)
add_page_break()
sections.section_2(doc, META, H)
add_page_break()
sections.section_3(doc, H)
add_page_break()
sections.section_4(doc, H)
add_page_break()
sections.section_5(doc, H)
add_page_break()
sections.section_6(doc, H)
add_page_break()
sections.section_7(doc, H)
add_page_break()
sections.section_8(doc, H)
add_page_break()
sections.section_9(doc, H)
add_page_break()
sections.section_10(doc, H)

# ---------------------------------------------------------------------------
out_path = os.path.join(ROOT, "BaoCao_Assignment02.docx")
doc.save(out_path)
print("Saved:", out_path)
print("Total figures inserted:", FIG_COUNTER["n"])
